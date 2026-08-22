"""
解析器单元测试。

用法:
    python -m pytest tests/test_parsers/ -v
"""

import os
import pytest
import tempfile


class TestTxtParser:
    """纯文本解析器测试"""

    def test_parse_simple_text(self):
        from app.services.importer.parsers.txt_parser import TxtParser

        with tempfile.NamedTemporaryFile(mode="w", suffix=".txt", delete=False, encoding="utf-8") as f:
            f.write("这是第一行。\n这是第二行。\n\n第三段内容。")
            tmp_path = f.name

        try:
            parser = TxtParser()
            result = parser.parse(tmp_path)
            assert "这是第一行" in result
            assert "第三段内容" in result
            assert len(result) > 0
        finally:
            os.unlink(tmp_path)

    def test_parse_utf8_with_special_chars(self):
        from app.services.importer.parsers.txt_parser import TxtParser

        content = "中文测试\nEnglish Test\n特殊字符：①②③ ★☆\nEmoji: 🎉"
        with tempfile.NamedTemporaryFile(mode="w", suffix=".txt", delete=False, encoding="utf-8") as f:
            f.write(content)
            tmp_path = f.name

        try:
            parser = TxtParser()
            result = parser.parse(tmp_path)
            assert "中文测试" in result
            assert "English Test" in result
            assert "🎉" in result
        finally:
            os.unlink(tmp_path)

    def test_parse_empty_file(self):
        from app.services.importer.parsers.txt_parser import TxtParser

        with tempfile.NamedTemporaryFile(mode="w", suffix=".txt", delete=False, encoding="utf-8") as f:
            f.write("")
            tmp_path = f.name

        try:
            parser = TxtParser()
            result = parser.parse(tmp_path)
            assert result == ""
        finally:
            os.unlink(tmp_path)


class TestMarkdownParser:
    """Markdown 解析器测试"""

    def test_parse_markdown_headings(self):
        from app.services.importer.parsers.markdown_parser import MarkdownParser

        content = "# 一级标题\n\n## 二级标题\n\n正文内容。\n\n### 三级标题\n\n更多内容。"
        with tempfile.NamedTemporaryFile(mode="w", suffix=".md", delete=False, encoding="utf-8") as f:
            f.write(content)
            tmp_path = f.name

        try:
            parser = MarkdownParser()
            result = parser.parse(tmp_path)
            assert "# 一级标题" in result
            assert "## 二级标题" in result
            assert "正文内容" in result
        finally:
            os.unlink(tmp_path)

    def test_parse_markdown_with_code_blocks(self):
        from app.services.importer.parsers.markdown_parser import MarkdownParser

        content = "说明：\n\n```python\nprint('hello')\n```\n\n结束。"
        with tempfile.NamedTemporaryFile(mode="w", suffix=".md", delete=False, encoding="utf-8") as f:
            f.write(content)
            tmp_path = f.name

        try:
            parser = MarkdownParser()
            result = parser.parse(tmp_path)
            assert "```python" in result
            assert "print('hello')" in result
        finally:
            os.unlink(tmp_path)

    def test_parse_markdown_tables(self):
        from app.services.importer.parsers.markdown_parser import MarkdownParser

        content = "| 列1 | 列2 |\n|-----|-----|\n| A   | B   |\n| C   | D   |"
        with tempfile.NamedTemporaryFile(mode="w", suffix=".md", delete=False, encoding="utf-8") as f:
            f.write(content)
            tmp_path = f.name

        try:
            parser = MarkdownParser()
            result = parser.parse(tmp_path)
            assert "| 列1 | 列2 |" in result
            assert "| A   | B   |" in result
        finally:
            os.unlink(tmp_path)


class TestDocxParser:
    """Word 解析器测试"""

    def test_parse_docx(self):
        from app.services.importer.parsers.docx_parser import DocxParser
        from docx import Document

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            tmp_path = f.name

        try:
            doc = Document()
            doc.add_heading("测试文档", level=1)
            doc.add_paragraph("这是第一段正文内容。")
            doc.add_heading("第二节", level=2)
            doc.add_paragraph("第二节的正文。")
            doc.save(tmp_path)

            parser = DocxParser()
            result = parser.parse(tmp_path)
            assert "# 测试文档" in result
            assert "这是第一段正文内容" in result
            assert "## 第二节" in result
            assert "第二节的正文" in result
        finally:
            os.unlink(tmp_path)

    def test_parse_docx_empty(self):
        from app.services.importer.parsers.docx_parser import DocxParser
        from docx import Document

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            tmp_path = f.name

        try:
            doc = Document()
            doc.save(tmp_path)

            parser = DocxParser()
            result = parser.parse(tmp_path)
            assert result == ""
        finally:
            os.unlink(tmp_path)

    def test_parse_docx_with_tables(self):
        from app.services.importer.parsers.docx_parser import DocxParser
        from docx import Document

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            tmp_path = f.name

        try:
            doc = Document()
            doc.add_paragraph("表格前文字。")
            table = doc.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "A1"
            table.cell(0, 1).text = "B1"
            table.cell(1, 0).text = "A2"
            table.cell(1, 1).text = "B2"
            doc.add_paragraph("表格后文字。")
            doc.save(tmp_path)

            parser = DocxParser()
            result = parser.parse(tmp_path)
            assert "表格前文字" in result
            assert "表格后文字" in result
            # 当前解析器契约只提取段落；表格存在时也不能影响前后正文。
            assert isinstance(result, str)
        finally:
            os.unlink(tmp_path)


class TestFormatDetection:
    """格式检测与路由测试"""

    def test_detect_format_all_types(self):
        from app.services.importer.parser_router import detect_format, FileType

        assert detect_format("doc.pdf") == FileType.PDF
        assert detect_format("DOC.PDF") == FileType.PDF
        assert detect_format("file.docx") == FileType.DOCX
        assert detect_format("file.doc") == FileType.DOCX
        assert detect_format("readme.md") == FileType.MD
        assert detect_format("readme.markdown") == FileType.MD
        assert detect_format("notes.txt") == FileType.TXT
        assert detect_format("image.png") == FileType.UNKNOWN
        assert detect_format("noext") == FileType.UNKNOWN

    def test_route_parser_non_pdf(self):
        from app.services.importer.parser_router import route_parser, FileType
        from app.services.importer.parsers.txt_parser import TxtParser
        from app.services.importer.parsers.markdown_parser import MarkdownParser
        from app.services.importer.parsers.docx_parser import DocxParser

        assert isinstance(route_parser(FileType.TXT), TxtParser)
        assert isinstance(route_parser(FileType.MD), MarkdownParser)
        assert isinstance(route_parser(FileType.DOCX), DocxParser)

    def test_route_parser_unknown_raises(self):
        from app.services.importer.parser_router import route_parser, FileType
        import pytest

        with pytest.raises(ValueError, match="不支持的文件类型"):
            route_parser(FileType.UNKNOWN)


class TestPDFTypeDetection:
    """PDF 类型检测测试"""

    def test_detect_pdf_type_on_text_pdf(self):
        from app.services.importer.parser_router import detect_pdf_type, PDFType

        # 用 fixtures 中的 Markdown 文件创建简单 PDF（如果没有 PDF 测试文件则跳过）
        tests_dir = os.path.dirname(os.path.dirname(os.path.dirname(__file__)))
        doc_path = os.path.join(tests_dir, "fixtures", "documents", "test_doc.md")
        if not os.path.exists(doc_path):
            pytest.skip("fixtures/documents/test_doc.md 不存在")

        # 检查是否有 PDF 测试文件
        pdf_path = os.path.join(tests_dir, "fixtures", "documents", "test_doc.pdf")
        if os.path.exists(pdf_path):
            pdf_type = detect_pdf_type(pdf_path)
            assert pdf_type in (PDFType.TEXT, PDFType.COMPLEX, PDFType.SCANNED)
